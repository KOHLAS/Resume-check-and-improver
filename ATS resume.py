import streamlit as st
import pdfplumber
from docx import Document
import spacy
from io import BytesIO

# Load the spaCy language model
nlp = spacy.load("en_core_web_sm")

# Function to extract text from PDF
def extract_pdf_text(file):
    with pdfplumber.open(file) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
        return text

# Function to extract text from DOCX
def extract_docx_text(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to check for required sections in the resume
def check_resume_structure(resume_text):
    structure_issues = []
    sections = ["work experience", "education", "skills"]
    
    for section in sections:
        if section not in resume_text.lower():
            structure_issues.append(f"Missing section: {section.title()}")
    
    return structure_issues

# Function to check for keyword matching between resume and job description
def check_keywords(resume_text, job_description_text):
    resume_doc = nlp(resume_text)
    job_desc_doc = nlp(job_description_text)

    resume_keywords = {token.lemma_.lower() for token in resume_doc if not token.is_stop}
    job_desc_keywords = {token.lemma_.lower() for token in job_desc_doc if not token.is_stop}

    matched_keywords = resume_keywords.intersection(job_desc_keywords)
    return matched_keywords, len(matched_keywords) / len(job_desc_keywords)

# Function to improve the resume by adding a Skills section
def add_skills_section(docx_file, skills_list):
    doc = Document(docx_file)
    
    # Add "Skills" section at the end of the document
    doc.add_heading("Skills", level=1)
    
    for skill in skills_list:
        doc.add_paragraph(skill)
    
    # Save to BytesIO to return the document in-memory
    output = BytesIO()
    doc.save(output)
    return output

# Streamlit App Interface
st.title("ATS Resume Checker and Improver")

# Step 1: Upload Resume
uploaded_resume = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])

# Step 2: Upload Job Description
uploaded_job_desc = st.file_uploader("Upload job description (Optional)", type=["pdf", "docx", "txt"])

# Step 3: Process the Resume
if st.button("Check Resume"):
    if uploaded_resume is not None:
        # Extract text from the uploaded resume
        if uploaded_resume.type == "application/pdf":
            resume_text = extract_pdf_text(uploaded_resume)
        elif uploaded_resume.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = extract_docx_text(uploaded_resume)

        st.write("### Resume Analysis")
        st.write(resume_text)

        # Perform ATS checks
        structure_issues = check_resume_structure(resume_text)
        st.write("### Structure Issues")
        if structure_issues:
            for issue in structure_issues:
                st.write(f"- {issue}")
        else:
            st.write("No structure issues detected!")

        # Keyword matching if a job description is provided
        if uploaded_job_desc is not None:
            if uploaded_job_desc.type == "application/pdf":
                job_description_text = extract_pdf_text(uploaded_job_desc)
            elif uploaded_job_desc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                job_description_text = extract_docx_text(uploaded_job_desc)
            elif uploaded_job_desc.type == "text/plain":
                job_description_text = uploaded_job_desc.read().decode("utf-8")

            matched_keywords, match_percentage = check_keywords(resume_text, job_description_text)
            st.write("### Keyword Match Results")
            st.write(f"Matched Keywords: {', '.join(matched_keywords)}")
            st.write(f"Match Percentage: {match_percentage * 100:.2f}%")

# Step 4: Allow Improvements in Word Document
if uploaded_resume is not None and uploaded_resume.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
    skills_to_add = st.text_area("Add Skills to the Resume (separate by commas)", value="Python, Project Management, Revit")

    if st.button("Improve Resume"):
        skills_list = [skill.strip() for skill in skills_to_add.split(",")]
        updated_resume = add_skills_section(uploaded_resume, skills_list)

        # Provide the user with a downloadable link for the improved resume
        st.download_button(
            label="Download Improved Resume",
            data=updated_resume.getvalue(),
            file_name="Improved_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
